from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import openai
from langdetect import detect


RESUME_SECTION_KEYWORDS = {
    "projets_assignés": ["projets assignés", "project assigned", "assigned projects"],
    "liste_projets": ["liste des projets", "project list", "geff", "gvc", "gts", "morseff"]
}

def extract_relevant_sections_from_resume(path):
    from docx import Document
    import re

    doc = Document(path)
    sections = {"projets_assignés": [], "liste_projets": []}
    paragraphs = doc.paragraphs
    tables = doc.tables

    # Track all tables with their index positions in doc._element.body
    table_elements = [t._element for t in tables]
    body_elements = list(doc._element.body)

    # Match section titles and extract only the first table *after* the paragraph
    for i, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        para_element = para._element

        for section_key, keywords in RESUME_SECTION_KEYWORDS.items():
            if any(re.search(r'\b' + re.escape(k) + r'\b', text) for k in keywords):

                # find first table after this paragraph in the XML tree
                try:
                    para_idx = body_elements.index(para_element)
                    for el in body_elements[para_idx+1:]:
                        if el in table_elements:
                            table = tables[table_elements.index(el)]
                            rows = [
                                " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                                for row in table.rows
                            ]
                            rows = [r for r in rows if r.strip()]
                            if rows:
                                sections[section_key].extend(rows)
                            break
                except Exception as e:
                    print(f"[ERROR] {e}")
                break

    return {k: "\n".join(v).strip() for k, v in sections.items() if v}


def match_resume_section(text):
    lower = text.lower()
    for section, keywords in RESUME_SECTION_KEYWORDS.items():
        if any(k in lower for k in keywords):
            return section
    return None

def filter_resume_section(section_name, content, ao_title):
    if not content.strip():
        return ""

    lang = detect(content)
    if lang == "fr":
        lang_instruction = "You must write your answer in French."
    else:
        lang_instruction = "You must write your answer in English."

    prompt = (
       f"You are adapting a candidate's {section_name} section from a CV.\n"
       f"{lang_instruction}\n"
       f"Only retain and slightly rephrase experience or projects that are clearly relevant to the AO titled: '{ao_title}'.\n"
       f"Remove unrelated parts. Don't fabricate.\n\n"
       f"Here is the original content:\n{content}\n\n"
       f"Return only the cleaned-up and filtered section content."
    )

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a CV editor aligning candidate experience to a specific opportunity."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.4
    )

    return response['choices'][0]['message']['content'].strip()

def adapt_resume_sections(raw_sections, ao_title):
    adapted = {}
    for section, content in raw_sections.items():
        filtered = filter_resume_section(section, content, ao_title)
        if "no content" not in filtered.lower():  # skip empty/irrelevant sections
            adapted[section] = filtered
    return adapted


def format_paragraph(p, bold=False):
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(11)
    run.bold = bold

def clear_paragraphs(doc):
    body = doc._element.body
    for child in list(body):
        body.remove(child)
def save_filtered_resume(original_path, output_path, adapted_sections):
    from docx import Document
    from docx.oxml import OxmlElement
    from shutil import copyfile
    import re

    if not adapted_sections:
        copyfile(original_path, output_path)
        return output_path

    doc = Document(original_path)
    body = doc._element.body
    paragraphs = doc.paragraphs
    tables = doc.tables

    for section_key in RESUME_SECTION_KEYWORDS:
        content = adapted_sections.get(section_key, "").strip()
        content_lower = content.lower()
        delete_entire = (
            not content
            or "does not have any relevant experience" in content_lower
            or "no content to return" in content_lower
        )

        for i, para in enumerate(paragraphs):
            text = para.text.strip().lower()
            if any(k in text for k in RESUME_SECTION_KEYWORDS[section_key]):
                p_el = para._element
                elements = list(body)
                try:
                    idx = elements.index(p_el)
                except ValueError:
                    continue

                if delete_entire:
                    p_el.getparent().remove(p_el)

                # Find the first table after this paragraph
                for el in elements[idx+1:]:
                    if el.tag.endswith("tbl"):
                        table = tables[[t._element for t in tables].index(el)]

                        if delete_entire:
                            el.getparent().remove(el)
                        else:
                            # FILTER each row based on whether its first line appears in GPT result
                            rows = table.rows[:]
                            for row in rows:
                                first_cell = row.cells[0].text.strip()
                                first_line = first_cell.splitlines()[0].strip()
                                if ":" in first_line:
                                    project_name = first_line.split(":", 1)[1].strip()
                                else:
                                    project_name = first_line.strip()
                                normalized_row_project = re.sub(r"[^a-z0-9]", "", project_name.lower())
                                normalized_gpt_projects = [
                                    re.sub(r"[^a-z0-9]", "", block.strip().lower())
                                    for block in content.split("\n\n") if block.strip() ]
                                keep = any(normalized_row_project in gpt_block for gpt_block in normalized_gpt_projects)

                                if not keep:
                                    row._element.getparent().remove(row._element)
                        break
                break

    if hasattr(output_path, "write"):
        doc.save(output_path)
    else:
        doc.save(str(output_path))

    return output_path


def extract_projects_from_tables(doc):
    tables = [table for table in doc.tables if any("Projet" in cell.text for row in table.rows for cell in row.cells)]
    projects = []
    for table in tables:
        for row in table.rows[1:]:  # skip header
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 4:
                projects.append({
                    "titre": cells[0],
                    "annee": cells[1],
                    "secteur": cells[2],
                    "description": cells[3],
                })
    return projects

def filter_projects_by_ao(projects, ao_title):
    import openai
    filtered = []
    for project in projects:
        prompt = (
            f"Below is a project entry from a CV.\n"
            f"Project title: {project['titre']}\n"
            f"Sector: {project['secteur']}\n"
            f"Description: {project['description']}\n"
            f"\nDetermine if this project is relevant to the AO titled: '{ao_title}'.\n"
            f"If yes, return a cleaned version of the description. If not, respond 'irrelevant'."
        )

        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are reviewing past project descriptions from a CV."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        result = response['choices'][0]['message']['content'].strip()
        if result.lower() != "irrelevant":
            project['description'] = result
            filtered.append(project)

    return filtered

def write_filtered_projects(doc, projects):
    if not projects:
        return
    p = doc.add_paragraph("Projets filtrés (tables)")
    p.runs[0].bold = True
    for proj in projects:
        doc.add_paragraph(f"- {proj['titre']} ({proj['annee']})")
        doc.add_paragraph(proj['description'])