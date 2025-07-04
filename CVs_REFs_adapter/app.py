# ✅ Combined adapter: references + resumes
import streamlit as st
import openai
import tempfile
import os
from pathlib import Path
from io import BytesIO

from .docx_adapter import get_sections_from_docx, save_adapted_docx
from .gpt_logic import adapt_all_sections
from .resume_adapter import (
    extract_relevant_sections_from_resume,
    adapt_resume_sections,
    save_filtered_resume,
    extract_projects_from_tables,
    filter_projects_by_ao,
    write_filtered_projects
)

openai.api_key = st.secrets["openai"]["api_key"]


@st.cache_data(show_spinner=False)
def cached_adaptation(sections_list, ao_title):
    return adapt_all_sections(sections_list, ao_title)

@st.cache_data(show_spinner=False)
def cached_resume_adaptation(raw_sections, ao_title):
    return adapt_resume_sections(raw_sections, ao_title)

if "ref_result" not in st.session_state:
    st.session_state["ref_result"] = None

if "resume_results" not in st.session_state:
    st.session_state["resume_results"] = []

def run_app():
    ao_title = st.text_input("AO Title", placeholder="Enter the AO title...")
    uploaded_ref = st.file_uploader("Upload Reference .docx file", type=["docx"], key="ref")
    uploaded_resumes = st.file_uploader("Upload Resumes (.docx)", type=["docx"], accept_multiple_files=True)

    # ✅ Reset outputs when input is cleared
    if not uploaded_ref:
        st.session_state["ref_result"] = None
    if not uploaded_resumes:
        st.session_state["resume_results"] = []


    if "submit_refs_cvs" not in st.session_state:
        st.session_state["submit_refs_cvs"] = False

    if st.button("Submit", key="submit_refs_cvs_btn"):
        st.session_state["submit_refs_cvs"] = True

    if st.session_state["submit_refs_cvs"] and ao_title and not st.session_state.get("already_processed", False):
        if uploaded_ref and ao_title:
            with st.spinner("Processing file(s)..."):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp.write(uploaded_ref.read())
                    input_path = tmp.name

                sections_list = get_sections_from_docx(input_path)
                if not sections_list:
                    st.error("❌ No sections extracted from reference file.")
                else:
                    adapted = cached_adaptation(sections_list, ao_title)
                    original_name = Path(uploaded_ref.name).stem
                    output_name = original_name + "_adapted.docx"
                    output_path = os.path.join(tempfile.gettempdir(), output_name)
                    save_adapted_docx(input_path, output_path, adapted, ao_title)


                    # ✅ Translate final output before showing download
                    try:
                        from docx import Document
                        from deep_translator import GoogleTranslator
                        from langdetect import detect

                        def translate_paragraph(para, target_lang="en"):
                            full_text = para.text.strip()
                            if full_text:
                                try:
                                    translated = GoogleTranslator(source='auto', target=target_lang).translate(full_text)
                                    for run in para.runs:
                                        run.text = ""
                                    para.runs[0].text = translated
                                except Exception as e:
                                    print("[Final Translation Error]", e)

                        ao_lang = detect(ao_title)
                        doc = Document(output_path)

                        for para in doc.paragraphs:
                            translate_paragraph(para, ao_lang)

                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        translate_paragraph(para, ao_lang)

                        doc.save(output_path)
                    except Exception as e:
                        print("[Final Translation Skipped]", e)

                    # ✅ Show download button
                    with open(output_path, "rb") as f:
                        st.session_state["ref_result"] = {
                            "data": f.read(),
                            "name": output_name,
                            "original": uploaded_ref.name
                        }




        # === Handle Resumes ===
        if uploaded_resumes and ao_title:
            for uploaded_resume in uploaded_resumes:
                with st.spinner(f"Processing {uploaded_resume.name}..."):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                        tmp.write(uploaded_resume.read())
                        resume_path = tmp.name

                    raw_sections = extract_relevant_sections_from_resume(resume_path)

                    adapted_resume = cached_resume_adaptation(raw_sections, ao_title)

                    original_name = Path(uploaded_resume.name).stem
                    output_name = original_name + "_adapted.docx"
                    output_buffer = BytesIO()
                    save_filtered_resume(resume_path, output_buffer, adapted_resume)
                    output_buffer.seek(0)
                    st.session_state["resume_results"].append({
                    "data": output_buffer.getvalue(),
                    "name": output_name,
                    "original": uploaded_resume.name
                })
        st.session_state["already_processed"] = True
    if st.session_state["ref_result"]:
        st.download_button(
            f"Download Adapted References: {st.session_state['ref_result']['original']}",
            data=st.session_state["ref_result"]["data"],
            file_name=st.session_state["ref_result"]["name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    for result in st.session_state["resume_results"]:
        st.download_button(
            f"Download Adapted Resume: {result['original']}",
            data=result["data"],
            file_name=result["name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )




