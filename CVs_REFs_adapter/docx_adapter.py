from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from langdetect import detect
from deep_translator import GoogleTranslator
from docx.text.paragraph import Paragraph


import os

SECTION_KEYWORDS = {
    "description du projet": ["description"],
    "services fournis": ["services"],
    "résultats issus du projet": ["résultats issus"],
    "projet et activités réalisée": ["activités", "projet"]
}
REFERENCE_TRIGGER = "référence du projet"

def match_section_title(text):
    lower_text = text.lower()
    for title, keywords in SECTION_KEYWORDS.items():
        if all(k in lower_text for k in keywords):
            return title
    return None

def format_paragraph(p, bold=False):
    run = p.runs[0] if p.runs else p.add_run()
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(11)
    run.bold = bold
    p.paragraph_format.alignment = 3  # Justify
    p.paragraph_format.right_indent = Cm(0.2)

def clear_cell(cell):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

def get_refs_tables(doc):
    return [table for table in doc.tables if any(REFERENCE_TRIGGER in cell.text.lower() for row in table.rows for cell in row.cells)]

def get_sections_from_docx(path):
    doc = Document(path)
    refs = []
    for table in get_refs_tables(doc):
        cells = [cell for row in table.rows for cell in row.cells if cell.text.strip()]
        lines = [line.strip() for cell in cells for line in cell.text.strip().split("\n") if line.strip()]
        sections = {}
        current_section = None
        for line in lines:
            matched = match_section_title(line)
            if matched:
                current_section = matched
                sections[current_section] = []
            elif current_section:
                sections[current_section].append(line)
        last = match_section_title(lines[-1]) if lines else None
        if last and last not in sections:
            sections[last] = []
        refs.append({k: "\n".join(v).strip() for k, v in sections.items()})
    return refs


def save_adapted_docx(input_path, output_path, adapted_refs, ao_title):
    doc = Document(input_path)
    tables = get_refs_tables(doc)

    for ref_index, (table, adapted) in enumerate(zip(tables, adapted_refs)):
        for row in table.rows:
            for cell in row.cells:
                found_desc_services = False
                found_resultats = False
                for para in list(cell.paragraphs):  # Use list to avoid runtime mutation
                    if not found_desc_services and any(
                        match_section_title(para.text.strip().lower()) in ["projet et activités réalisée", "description du projet"]
                        for _ in [0]
                    ):
                        clear_cell(cell)
                        for section in ["projet et activités réalisée", "description du projet", "services fournis"]:
                            if section in adapted:
                                p_title = cell.add_paragraph(section.capitalize())
                                format_paragraph(p_title, bold=True)
                                for line in adapted[section].split("\n"):
                                    if line.strip():
                                        p = cell.add_paragraph(line.strip())
                                        format_paragraph(p)
                        found_desc_services = True

                    elif not found_resultats and match_section_title(para.text.strip().lower()) == "résultats issus du projet":
                        if "résultats issus du projet" in adapted:
                            clear_cell(cell)
                            p_title = cell.add_paragraph("Résultats issus du projet")
                            format_paragraph(p_title, bold=True)
                            for line in adapted["résultats issus du projet"].split("\n"):
                                if line.strip():
                                    p = cell.add_paragraph(line.strip())
                                    format_paragraph(p)
                        found_resultats = True

    if hasattr(output_path, "write"):
        doc.save(output_path)
    else:
        doc.save(str(output_path))

    return output_path
