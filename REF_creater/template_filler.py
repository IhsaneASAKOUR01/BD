from docx import Document
from .gpt_extract import extract_field, force_field_completion
import time
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


from docx.oxml import OxmlElement

# Fields GPT will extract
GPT_FIELDS = [
    "Catégorie de service",
    "Nom de la mission",
    "Pays",
    "Localisation dans le pays",
    "Nom du client",
    "Description du projet",
    "Services fournis",
    "Résultats issus du projet"
]

# Fields user fills manually
USER_FIELDS = [
    "Référence du projet",
    "Valeur approximative du contrat (en dollars US courants)",
    "Durée de la mission (en mois)",
    "Nombre total de mois-personnes pour la mission",
    "Personne à contacter, titre/fonction, téléphone/adresse",
    "Date de début (mois/année)",
    "Date d'achèvement (mois/année)",
    "Nombre de mois-personnes professionnels fournis par votre cabinet/organisation ou vos sous-traitants",
    "Nom des consultants associés, le cas échéant",
    "Nom du personnel professionnel senior de votre cabinet/organisation impliqué et fonction et/ou rôles exercés (par exemple, directeur/coordonateur de projet, chef d'équipe)"
]

def fill_template(template_path, report_text):
    doc = Document(template_path)
    for para in doc.paragraphs:
        for field in GPT_FIELDS:
            if para.text.strip().startswith(field):
                value = extract_field(field, report_text)
                if value:
                    para.text = f"{field} : {value}"
    return doc

def fill_template_with_debug(template_path, report_text):
    from .gpt_extract import extract_field
    import time

    GPT_FIELDS = [
        "Catégorie de service",
        "Nom de la mission",
        "Pays",
        "Localisation dans le pays",
        "Nom du client",
        "Description du projet",
        "Services fournis",
        "Résultats issus du projet"
    ]

    results = extract_field(report_text, GPT_FIELDS)
    results = force_field_completion(results)
    return results


def fill_template_with_values(template_path, extracted_fields):
    doc = Document(template_path)
    for para in doc.paragraphs:
        for field in GPT_FIELDS:
            if para.text.strip().startswith(field):
                value = extracted_fields.get(field)
                if value:
                    para.text = f"{field} : {value}"
    return doc

def clear_cell(cell):
    cell._element.clear_content()


def fill_reference_table(template_path, extracted_fields):
    doc = Document(template_path)

    def make_bullet_paragraph(cell, text):
        p = cell.add_paragraph()
        p.style = doc.styles['Normal']
        p_paragraph = p._p

        # Set numbering properties (Word bullet level 0)
        numPr = OxmlElement('w:numPr')

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)

        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(numId)

        p_props = p_paragraph.get_or_add_pPr()
        p_props.append(numPr)

        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(11)

        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field in extracted_fields:
                    if cell.text.strip().lower().startswith(field.lower()):
                        if extracted_fields[field]:
                            clear_cell(cell)

                            field_text = f"{field} :"
                            value_text = extracted_fields[field]

                            # Title paragraph
                            p_title = cell.add_paragraph()
                            run_title = p_title.add_run(field_text)
                            run_title.bold = True
                            run_title.font.name = 'Times New Roman'
                            run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            run_title.font.size = Pt(11)
                            p_title.paragraph_format.space_before = Pt(6)
                            p_title.paragraph_format.space_after = Pt(6)

                            # Bullet formatting for specific fields
                            if field.lower() in ["services fournis", "résultats issus du projet"] and any(b in value_text for b in ["•", "-", "\n"]):
                                bullet_lines = [line.strip() for line in value_text.splitlines() if line.strip()]
                                for line in bullet_lines:
                                    if line.strip().endswith(":"):
                                        p_intro = cell.add_paragraph()
                                        run_intro = p_intro.add_run(line.strip("•- \t"))
                                        run_intro.font.name = 'Times New Roman'
                                        run_intro._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                                        run_intro.font.size = Pt(11)
                                        p_intro.paragraph_format.space_before = Pt(3)
                                        p_intro.paragraph_format.space_after = Pt(3)
                                    else:
                                        make_bullet_paragraph(cell, line.strip("•- \t"))

                            else:
                                # Normal value paragraph
                                p_value = cell.add_paragraph()
                                run_value = p_value.add_run(value_text)
                                run_value.bold = False
                                run_value.font.name = 'Times New Roman'
                                run_value._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                                run_value.font.size = Pt(11)

                            # Apply 0.2 cm margin to all sides
                            tc_pr = cell._tc.get_or_add_tcPr()
                            tcMar = OxmlElement('w:tcMar')
                            for side in ['top', 'left', 'bottom', 'right']:
                                mar = OxmlElement(f'w:{side}')
                                mar.set(qn('w:w'), '113')  # 0.2 cm = ~113 twips
                                mar.set(qn('w:type'), 'dxa')
                                tcMar.append(mar)
                            tc_pr.append(tcMar)

    return doc
