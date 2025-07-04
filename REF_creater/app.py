import streamlit as st
from pathlib import Path
import tempfile
from docx import Document
from .gpt_extract import extract_field
from .template_filler import fill_reference_table, fill_template_with_debug
from .utils import load_report_text
from deep_translator import GoogleTranslator

st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Poppins:wght@300;600&display=swap" rel="stylesheet">

<style>
/* Fix layout for consistency */
.main .block-container {
    max-width: 900px !important;
    margin: auto !important;
    padding-left: 1rem !important;
    padding-right: 1rem !important;
}
.custom-title {
    text-align: center !important;
    font-size: 3rem !important;
    font-weight: 600 !important;
    color: #ffffff !important;
    background: linear-gradient(135deg, #1d3557, #457b9d) !important;
    padding: 1.5rem !important;
    border-radius: 12px !important;
    margin-bottom: 1rem !important;
    box-shadow: 0 4px 10px rgba(0,0,0,0.1) !important;
}

.custom-desc {
    text-align: center !important;
    color: #333 !important;
    font-size: 1.1rem !important;
    margin-bottom: 2rem !important;
}

section[data-testid="stFileUploader"] {
    background: #ffffff !important;
    border: 2px dashed #339af0 !important;
    padding: 1.5rem !important;
    border-radius: 15px !important;
    box-shadow: 0 4px 12px rgba(0,0,0,0.05) !important;
}

button[kind="primary"] {
    background: linear-gradient(135deg, #339af0, #1c7ed6) !important;
    border: none !important;
    color: white !important;
    padding: 0.75rem 1.5rem !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    font-size: 1rem !important;
    transition: all 0.2s ease-in-out !important;
    margin-top: 1rem !important;
}
button[kind="primary"]:hover {
    background: linear-gradient(135deg, #1c7ed6, #339af0) !important;
    box-shadow: 0 0 10px rgba(51, 154, 240, 0.5) !important;
    transform: scale(1.03) !important;
}

.stDownloadButton > button {
    background-color: #339af0 !important;
    color: white !important;
    padding: 0.6rem 1.2rem !important;
    border-radius: 8px !important;
    font-weight: bold !important;
    border: none !important;
    margin: 1rem 0.5rem 0 0 !important;
    transition: 0.2s ease-in-out !important;
}
.stDownloadButton > button:hover {
    background-color: #1c7ed6 !important;
    transform: scale(1.05) !important;
}
</style>
""", unsafe_allow_html=True)


def run_app():
    if "field_values" not in st.session_state:
        st.session_state["field_values"] = None
    if "output_path_fr" not in st.session_state:
        st.session_state["output_path_fr"] = None
    if "output_path_en" not in st.session_state:
        st.session_state["output_path_en"] = None

    uploaded_report = st.file_uploader(
        "Upload Project Report (.docx, .pdf, .pptx, .txt)", 
        type=["docx", "pdf", "pptx", "txt"]
    )
    submit = st.button("Submit", key="submit_ref_creator")

    # ✅ Reset the generation if a different file is uploaded
    if uploaded_report and "last_uploaded" in st.session_state:
        if uploaded_report.name != st.session_state["last_uploaded"]:
            st.session_state["generated"] = False

    # ✅ Update the uploaded filename in session
    st.session_state["last_uploaded"] = uploaded_report.name if uploaded_report else None

    TEMPLATE_PATH = "REF_creater/Référence Template.docx"

    def translate_docx(input_path, output_path, target_lang="en"):
        doc = Document(input_path)
        
        def translate_paragraph(para):
            text = para.text.strip()
            if text:
                try:
                    translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
                    # Replace text of first run only
                    if para.runs:
                        para.runs[0].text = translated
                        for i in range(1, len(para.runs)):
                            para.runs[i].text = ""
                except Exception as e:
                    print("[Translation Error]", e)

        for para in doc.paragraphs:
            translate_paragraph(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        translate_paragraph(para)

        doc.save(output_path)

    # Process only once per upload
    if submit and uploaded_report:
        report_text = load_report_text(uploaded_report)

        with st.spinner("Extracting fields with GPT..."):
            st.session_state["field_values"] = fill_template_with_debug(TEMPLATE_PATH, report_text)

        with st.spinner("Generating French reference..."):
            filled_doc = fill_reference_table(TEMPLATE_PATH, st.session_state["field_values"])

            mission_name = st.session_state["field_values"].get("Nom de la mission", "output").strip()
            mission_name_safe = "".join(c for c in mission_name if c.isalnum() or c in (" ", "_", "-")).rstrip()

            # Save French version
            fr_name = f"{mission_name_safe}_ref_VF.docx"
            path_fr = Path(tempfile.gettempdir()) / fr_name
            filled_doc.save(path_fr)
            st.session_state["output_path_fr"] = path_fr

        with st.spinner("Translating to English..."):
            en_name = f"{mission_name_safe}_ref_VA.docx"
            path_en = Path(tempfile.gettempdir()) / en_name
            translate_docx(path_fr, path_en, target_lang="en")
            st.session_state["output_path_en"] = path_en

        st.session_state["generated"] = True

    # Download buttons
    if st.session_state["output_path_fr"]:
        with open(st.session_state["output_path_fr"], "rb") as f:
            st.download_button(
                label="📥 Download French Version",
                data=f,
                file_name=st.session_state["output_path_fr"].name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    if st.session_state["output_path_en"]:
        with open(st.session_state["output_path_en"], "rb") as f:
            st.download_button(
                label="📥 Download English Version",
                data=f,
                file_name=st.session_state["output_path_en"].name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
