import re
from docx import Document

EXACT_TITLES = ["Certification"]
FLEXIBLE_TITLES = [
    # French
    "Nom du personnel",
    "Poste proposé",
    "Employeur",
    "Date de naissance",
    "Nationalité",
    "Education",
    "Certifications professionnelles ou adhésions à des associations professionnelles",
    "Compétences",
    "Autres formations pertinentes",
    "Pays d'expérience professionnelle",
    "Langues",
    "Historique d'emploi",
    "Projets assignés",
    "Liste des projets",
    # English
    "Name of Staff",
    "Proposed Position",
    "Employer",
    "Date of Birth",
    "Nationality",
    "Education",
    "Professional Certification or Membership in Professional Associations",
    "Skills",
    "Other Relevant Training",
    "Countries of Work Experience",
    "Languages",
    "Employment Record",
    "Project assigned",
    "List of projects"
]

FLEXIBLE_TITLES = [t.lower() for t in FLEXIBLE_TITLES]

def is_exact_title(text):
    return text.strip() in EXACT_TITLES

def is_flexible_title(text):
    return any(text.lower().startswith(t) for t in FLEXIBLE_TITLES)

def split_paragraph_by_titles(text):
    matches = []
    lower_text = text.lower()
    for title in FLEXIBLE_TITLES:
        idx = lower_text.find(title)
        if idx != -1:
            matches.append((idx, title))
    matches.sort()
    if not matches:
        return [text]
    sections = []
    for i, (start, _) in enumerate(matches):
        end = matches[i + 1][0] if i + 1 < len(matches) else len(text)
        segment = text[start:end].strip()
        if segment:
            sections.append(segment)
    return sections

def extract_all_parts_by_section_titles(doc_path):
    doc = Document(doc_path)
    sections = []
    current = []
    section_idx = -1
    para_to_section = {}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        segments = split_paragraph_by_titles(text)
        for seg in segments:
            if is_exact_title(seg) or is_flexible_title(seg):
                if current:
                    sections.append("\n".join(current).strip())
                    current.clear()
                section_idx += 1
            current.append(seg)
            para_to_section[i] = section_idx

    if current:
        sections.append("\n".join(current).strip())

    para_iter = iter(enumerate(doc.paragraphs))
    table_idx = 0
    last_para_idx = -1

    for block in doc.element.body:
        if block.tag.endswith("p"):
            try:
                last_para_idx, para = next(para_iter)
                if not para.text.strip():
                    continue
            except StopIteration:
                continue
        elif block.tag.endswith("tbl"):
            lookup_idx = last_para_idx
            while lookup_idx not in para_to_section and lookup_idx >= 0:
                lookup_idx -= 1
            section_id = para_to_section.get(lookup_idx, -1)
            if 0 <= section_id < len(sections):
                table = doc.tables[table_idx]
                table_idx += 1
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    rows.append(" | ".join(cells))
                sections[section_id] += "\n" + "\n".join(rows)

    return sections

def extract_full_text(doc_path):
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(doc_path)
    text_parts = []

    def iter_block_items(parent):
        for child in parent.element.body:
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            txt = block.text.strip()
            if txt:
                text_parts.append(txt)
        elif isinstance(block, Table):
            rows = []
            for row in block.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                text_parts.append("\n".join(rows))

    return "\n\n".join(text_parts)
