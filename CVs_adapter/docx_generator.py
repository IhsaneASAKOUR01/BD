from docx import Document
from docx.shared import Pt
from difflib import SequenceMatcher
import re

def similar(a, b):
    return SequenceMatcher(None, a.lower().strip(), b.lower().strip()).ratio() > 0.75

def apply_style(run, bold=False):
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold

def find_exact_or_similar(label, filled_dict):
    label_clean = clean_label(label)
    for key in filled_dict.keys():
        if similar(label_clean, key):
            return filled_dict[key]
    return None

def clean_label(text):
    return re.sub(r"\{.*?\}", "", text).replace(":", "").strip()

def format_table_as_text(table_data):
    col_widths = [max(len(str(cell)) for cell in col) for col in zip(*table_data)]
    lines = [" | ".join(str(cell).ljust(col_widths[i]) for i, cell in enumerate(row)) for row in table_data]
    return "\n".join(lines)

def fill_docx_template_by_labels(template_path, filled_dict, output_path):
    doc = Document(template_path)

    used_labels = set()

    for para in doc.paragraphs:
        para_label = clean_label(para.text)
        value = find_exact_or_similar(para_label, filled_dict)

        if value and para_label not in used_labels:
            used_labels.add(para_label)
            if isinstance(value, str):
                para.text = f"{para_label}: {value}" if ":" not in para_label else value
                for run in para.runs:
                    apply_style(run)
            elif isinstance(value, list):
                para.text = para_label
                formatted_table = format_table_as_text(value)
                para.add_run("\n" + formatted_table)

    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        table_label = " | ".join(headers)
        value = find_exact_or_similar(table_label, filled_dict)

        if isinstance(value, list):
            while len(table.rows) > 1:
                tbl = table._tbl
                tbl.remove(tbl.tr_lst[1])
            for row_data in value[1:]:
                row = table.add_row()
                for idx, cell_text in enumerate(row_data):
                    if idx < len(row.cells):
                        row.cells[idx].text = cell_text
                        for para in row.cells[idx].paragraphs:
                            for run in para.runs:
                                apply_style(run)

    doc.save(output_path)
